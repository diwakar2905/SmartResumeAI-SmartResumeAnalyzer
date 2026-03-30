import fitz  # PyMuPDF
import re
from docx import Document
import spacy
import os
import logging

logger = logging.getLogger(__name__)

# Load the spaCy model - spaCy model ko load karte hain
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    logger.error("Spacy model not found. Please run: python -m spacy download en_core_web_sm")
    nlp = None


def extract_text_from_pdf(path):
    """Extract text from a PDF file using PyMuPDF (fitz) with multiple fallback methods - PDF se text extract karte hain multiple methods se."""
    text = ""
    try:
        with fitz.open(path) as doc:  # Use context manager for safe file handling - Safe file handling ke liye context manager use karte hain
            logger.info(f"PDF opened successfully. Pages: {len(doc)}")
            
            for page_num, page in enumerate(doc):
                logger.info(f"Processing page {page_num + 1}")
                
                # Method 1: Try standard text extraction - Method 1: Standard text extraction try karte hain
                try:
                    page_text = page.get_text("text")
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        logger.info(f"Page {page_num + 1}: Extracted {len(page_text)} characters")
                        continue
                except Exception as e:
                    logger.warning(f"Standard text extraction failed for page {page_num + 1}: {e}")
                
                # Method 2: Try HTML extraction and clean it - Method 2: HTML extraction try karte hain aur clean karte hain
                try:
                    html_text = page.get_text("html")
                    if html_text:
                        # Clean HTML tags - HTML tags ko clean karte hain
                        import re
                        clean_text = re.sub(r'<[^>]+>', '', html_text)
                        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                        if clean_text:
                            text += clean_text + "\n"
                            logger.info(f"Page {page_num + 1}: Extracted {len(clean_text)} characters via HTML")
                            continue
                except Exception as e:
                    logger.warning(f"HTML extraction failed for page {page_num + 1}: {e}")
                
                # Method 3: Try raw text extraction - Method 3: Raw text extraction try karte hain
                try:
                    raw_text = page.get_text("raw")
                    if raw_text and raw_text.strip():
                        text += raw_text + "\n"
                        logger.info(f"Page {page_num + 1}: Extracted {len(raw_text)} characters via raw")
                        continue
                except Exception as e:
                    logger.warning(f"Raw text extraction failed for page {page_num + 1}: {e}")
                
                # Method 4: Try to extract text from blocks - Method 4: Blocks se text extract karne ki koshish karte hain
                try:
                    blocks = page.get_text("dict")
                    page_text = ""
                    for block in blocks.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line.get("spans", []):
                                    page_text += span.get("text", "") + " "
                    if page_text.strip():
                        text += page_text + "\n"
                        logger.info(f"Page {page_num + 1}: Extracted {len(page_text)} characters via blocks")
                except Exception as e:
                    logger.warning(f"Block extraction failed for page {page_num + 1}: {e}")
                
                # Extract tables if present - Agar tables hain to extract karte hain
                try:
                    tables = page.get_tables()
                    for table in tables:
                        for row in table:
                            row_text = " | ".join([cell.strip() for cell in row if cell.strip()])
                            if row_text:
                                text += row_text + "\n"
                except Exception as e:
                    logger.warning(f"Table extraction failed for page {page_num + 1}: {e}")
                            
    except Exception as e:
        logger.error(f"Error reading PDF {path}: {e}")
        return None
    
    final_text = text.strip()
    logger.info(f"Total extracted text length: {len(final_text)} characters")
    
    if not final_text:
        logger.warning("Warning: No text extracted from PDF")
        return None
    
    return final_text


def extract_text_from_docx(path):
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = Document(path)
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error reading DOCX {path}: {e}")
        return None


def extract_basic_info(text):
    """Extract name, email, and phone number from the text using NLP and regex."""
    if not text:
        return {"name": None, "email": None, "phone": None}

    name = None
    email = None
    phone = None

    # Email extraction with improved regex
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        email = email_match.group()

    # Phone extraction with improved regex for various formats
    phone_patterns = [
        r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',  # US format
        r'\+?[0-9]{1,3}[-.\s]?[0-9]{1,4}[-.\s]?[0-9]{1,4}[-.\s]?[0-9]{1,4}',  # International
        r'\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',  # Simple US format
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group()
            break

    # Name extraction with multiple strategies
    name = extract_name_advanced(text)

    return {
        "name": name.strip() if name else "Not Found",
        "email": email if email else "Not Found",
        "phone": phone if phone else "Not Found"
    }


def extract_name_advanced(text):
    """Advanced name extraction using multiple strategies."""
    lines = text.split('\n')
    
    # Strategy 1: NLP-based extraction
    if nlp:
        try:
            doc = nlp(text)
            for ent in doc.ents:
                if ent.label_ == 'PERSON':
                    # Clean up the name
                    name = ent.text.strip()
                    if len(name.split()) >= 2 and len(name) < 50:  # Reasonable name length
                        return name
        except Exception as e:
            logger.warning(f"NLP name extraction failed: {e}")

    # Strategy 2: Look for capitalized name patterns in the first few lines
    name_patterns = [
        r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}$',  # First Last or First Middle Last
        r'^[A-Z][a-z]+\s+[A-Z]\.?\s*[A-Z][a-z]+$',  # First M. Last
        r'^[A-Z][a-z]+\s+[A-Z][a-z]+$',  # First Last
    ]
    
    # Check first 10 lines for name patterns
    for line in lines[:10]:
        line = line.strip()
        if not line or len(line) < 3:
            continue
            
        # Skip lines that are likely not names
        if any(skip_word in line.lower() for skip_word in ['resume', 'cv', 'phone', 'email', '@', 'www', 'http']):
            continue
            
        for pattern in name_patterns:
            match = re.match(pattern, line)
            if match:
                name = match.group()
                # Additional validation
                if len(name.split()) >= 2 and len(name) < 50:
                    return name

    # Strategy 3: Look for the most prominent capitalized line
    for line in lines[:15]:
        line = line.strip()
        if not line or len(line) < 3:
            continue
            
        # Skip obvious non-name lines
        if any(skip_word in line.lower() for skip_word in ['resume', 'cv', 'phone', 'email', '@', 'www', 'http', 'objective', 'summary']):
            continue
            
        # Check if line looks like a name (2-4 capitalized words)
        words = line.split()
        if 2 <= len(words) <= 4:
            # Check if most words start with capital letters
            capitalized_words = sum(1 for word in words if word and word[0].isupper())
            if capitalized_words >= len(words) * 0.8:  # 80% of words capitalized
                return line

    return None


def extract_contact_info(text):
    """Extract additional contact information."""
    contact_info = {}
    
    # LinkedIn
    linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/company/)[a-zA-Z0-9-]+'
    linkedin_match = re.search(linkedin_pattern, text)
    if linkedin_match:
        contact_info['linkedin'] = linkedin_match.group()
    
    # GitHub
    github_pattern = r'github\.com/[a-zA-Z0-9-]+'
    github_match = re.search(github_pattern, text)
    if github_match:
        contact_info['github'] = github_match.group()
    
    # Website
    website_pattern = r'(?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.[a-zA-Z]{2,}(?:/\S*)?'
    website_matches = re.findall(website_pattern, text)
    if website_matches:
        # Filter out common non-personal websites
        personal_sites = [site for site in website_matches 
                         if not any(exclude in site.lower() 
                                   for exclude in ['linkedin.com', 'github.com', 'gmail.com', 'yahoo.com'])]
        if personal_sites:
            contact_info['website'] = personal_sites[0]
    
    return contact_info


def extract_education_info(text):
    """Extract education information from resume text."""
    education_info = []
    
    # Common education keywords
    education_keywords = [
        r'education', r'degree', r'bachelor', r'master', r'phd', r'doctorate',
        r'university', r'college', r'school', r'institute', r'academy'
    ]
    
    lines = text.split('\n')
    in_education_section = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if we're entering education section
        if any(re.search(keyword, line, re.IGNORECASE) for keyword in education_keywords):
            in_education_section = True
            continue
            
        # If we're in education section, look for degree patterns
        if in_education_section:
            # Look for degree patterns
            degree_patterns = [
                r'(?:Bachelor|Master|PhD|Doctorate|Associate)\s+(?:of|in)\s+[A-Za-z\s]+',
                r'[A-Za-z]+\s+(?:University|College|Institute|School)',
                r'\d{4}',  # Year
            ]
            
            # If line contains multiple education indicators, it might be education info
            indicators = sum(1 for pattern in degree_patterns if re.search(pattern, line, re.IGNORECASE))
            if indicators >= 2:
                education_info.append(line)
    
    return education_info


def get_text_statistics(text):
    """Get basic statistics about the resume text."""
    if not text:
        return {}
    
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    
    return {
        'word_count': len(words),
        'sentence_count': len([s for s in sentences if s.strip()]),
        'character_count': len(text),
        'average_words_per_sentence': len(words) / len([s for s in sentences if s.strip()]) if sentences else 0,
        'has_email': bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)),
        'has_phone': bool(re.search(r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}', text)),
        'has_links': bool(re.search(r'https?://', text)),
    }
